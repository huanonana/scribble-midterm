from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent / "Bao_cao_Scribble_Phan_tan.docx"
BLUE = RGBColor(31, 78, 121)
INK = RGBColor(28, 37, 48)
MUTED = RGBColor(90, 100, 112)
LIGHT = "EAF1F8"


def set_font(run, size=11, bold=False, color=INK, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 18, BLUE),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11.5, RGBColor(49, 93, 135)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "SCRIBBLE DISTRIBUTED | BÁO CÁO BÀI TẬP LỚN"
    set_font(header.runs[0], size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, size=18, bold=True, color=BLUE)
    if subtitle:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(10)
        set_font(q.add_run(subtitle), size=10.5, color=MUTED)


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(item))


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    shade_paragraph(p, "F3F5F7")
    set_font(p.add_run(text), size=8.5, color=RGBColor(35, 45, 55), name="Consolas")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = header
        set_font(cell.paragraphs[0].runs[0], size=9, bold=True, color=BLUE)
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = value
            set_font(cells[i].paragraphs[0].runs[0], size=8.8)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return t


def page_break(doc):
    doc.add_page_break()


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("[TÊN TRƯỜNG / KHOA]"), size=12, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("BÁO CÁO BÀI TẬP LỚN"), size=18, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("PHÁT TRIỂN HỆ THỐNG SCRIBBLE CÓ KHẢ NĂNG XỬ LÝ PHÂN TÁN"), size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Replication qua HTTP và đồng bộ snapshot giữa các node"), size=12, color=MUTED)
    for _ in range(4):
        doc.add_paragraph()
    table(doc, ["Thông tin", "Nội dung"], [
        ("Môn học", "[Điền tên môn học]"),
        ("Giảng viên", "[Điền tên giảng viên]"),
        ("Nhóm", "[Điền số nhóm]"),
        ("Thành viên", "[Điền họ tên - MSSV của từng thành viên]"),
        ("Năm học", "2025 - 2026"),
    ], [1.6, 4.7])


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    page_break(doc)

    title(doc, "TÓM TẮT", "Mục tiêu, phạm vi và kết quả chính")
    para(doc, "Bài tập lựa chọn Scribble, một cơ sở dữ liệu JSON nhỏ gọn viết bằng Go, làm nền tảng nghiên cứu và thực nghiệm. Dự án gốc cung cấp các thao tác tạo cơ sở dữ liệu, ghi, đọc, đọc toàn bộ collection và xóa dữ liệu trên hệ thống tệp cục bộ.")
    para(doc, "Nhóm mở rộng Scribble theo hướng xử lý phân tán bằng hai tính năng mới: (1) sao chép thao tác ghi/xóa sang các peer qua HTTP; (2) tạo snapshot và đồng bộ dữ liệu để khởi tạo hoặc phục hồi một node. Hai tính năng được tích hợp trực tiếp vào Driver, có HTTP handler, chương trình demo nhiều node và kiểm thử tự động bằng httptest.")
    bullets(doc, [
        "Kết quả: toàn bộ test gốc và test mở rộng chạy thành công với go test ./....",
        "Phạm vi: mô hình replication đơn giản, ưu tiên khả năng học tập và quan sát luồng dữ liệu.",
        "Giới hạn: chưa có consensus, quorum, vector clock hoặc tự động xử lý xung đột.",
    ])
    table(doc, ["Hạng mục", "Kết quả"], [
        ("Mã nguồn", "Go module, distributed.go, distributed_test.go, demo node"),
        ("Tính năng 1", "HTTP peer replication cho Write/Delete"),
        ("Tính năng 2", "Snapshot endpoint và SyncFrom"),
        ("Quản lý mã nguồn", "Nhiều commit theo từng giai đoạn"),
    ], [1.8, 4.5])
    page_break(doc)

    title(doc, "MỤC LỤC VÀ CẤU TRÚC BÁO CÁO")
    table(doc, ["Trang", "Nội dung"], [
        ("1-2", "Bìa, tóm tắt"),
        ("3-5", "Dự án Scribble, mục đích, chức năng, ứng dụng"),
        ("6-8", "Cài đặt, cấu trúc mã nguồn, baseline"),
        ("9-13", "Thiết kế và hiện thực hai tính năng phân tán"),
        ("14-16", "Thực nghiệm, kết quả, đánh giá"),
        ("17-19", "GitHub, đóng góp, kết luận và tài liệu tham khảo"),
    ], [1.2, 5.1])
    para(doc, "Báo cáo tập trung vào những thay đổi có thể kiểm chứng trong mã nguồn. Các lệnh, endpoint và kết quả test được ghi lại để người đọc có thể tái tạo thực nghiệm trên máy khác.")
    bullets(doc, [
        "Repo gốc: https://github.com/sdomino/scribble",
        "Ngôn ngữ: Go; dữ liệu lưu dưới dạng từng tệp JSON.",
        "Môi trường thực nghiệm: Windows PowerShell, Go module.",
    ])
    page_break(doc)

    title(doc, "1. GIỚI THIỆU DỰ ÁN SCRIBBLE")
    para(doc, "Scribble là một thư viện cơ sở dữ liệu JSON nhúng. Thay vì chạy một tiến trình database server riêng, ứng dụng gọi trực tiếp thư viện và dữ liệu được ghi xuống thư mục do người dùng chỉ định. Mỗi collection là một thư mục; mỗi resource là một tệp JSON.")
    para(doc, "Thiết kế tối giản giúp Scribble phù hợp cho ứng dụng nhỏ, prototype, công cụ dòng lệnh và bài thực hành về lưu trữ. Tuy nhiên, phiên bản gốc chỉ hoạt động trên một máy và không có cơ chế giao tiếp giữa các node.")
    table(doc, ["Thuộc tính", "Mô tả"], [
        ("Kiểu hệ thống", "Embedded JSON database"),
        ("Đơn vị dữ liệu", "Collection / Resource"),
        ("Cơ chế lưu", "Tệp JSON + rename file tạm"),
        ("Đồng thời", "Mutex theo collection"),
        ("Phân tán bản gốc", "Không hỗ trợ"),
    ], [1.7, 4.6])
    page_break(doc)

    title(doc, "2. MỤC ĐÍCH, CHỨC NĂNG VÀ ỨNG DỤNG")
    para(doc, "Mục đích chính của Scribble là cung cấp API lưu trữ đơn giản, dễ tích hợp và không cần vận hành dịch vụ ngoài. API cốt lõi gồm New, Write, Read, ReadAll và Delete.")
    bullets(doc, [
        "Ứng dụng cấu hình cục bộ, cache bền vững, dữ liệu thử nghiệm.",
        "Prototype cần lưu JSON nhanh mà không cần cài PostgreSQL hoặc MongoDB.",
        "Công cụ học tập về file I/O, JSON serialization, mutex và atomic rename.",
        "Sau mở rộng: demo replication, bootstrap node và phục hồi dữ liệu.",
    ])
    para(doc, "Trong thực tế, mô hình này không thay thế cơ sở dữ liệu phân tán hoàn chỉnh. Giá trị của bài tập nằm ở việc biến một thư viện đơn node thành hệ thống có giao tiếp node-to-node, quan sát được lỗi mạng và đánh giá trade-off nhất quán.")
    page_break(doc)

    title(doc, "3. PHÂN TÍCH KIẾN TRÚC PHIÊN BẢN GỐC")
    code(doc, "Application -> Driver.Write(collection, resource, value)\n            -> mutex theo collection\n            -> json.MarshalIndent\n            -> ghi resource.json.tmp\n            -> os.Rename thành resource.json")
    para(doc, "Driver quản lý thư mục dữ liệu, logger và bản đồ mutex. Mutex toàn cục bảo vệ quá trình tạo mutex cho từng collection; mutex collection ngăn ghi/xóa đồng thời vào cùng nhóm dữ liệu.")
    bullets(doc, [
        "Ưu điểm: code ngắn, dễ đọc; rename file tạm giảm nguy cơ tệp bị ghi dở.",
        "Nhược điểm: không có WAL, version, network protocol hoặc metadata node.",
        "Điểm mở rộng phù hợp: sau thao tác local thành công, tạo operation và gửi sang peer.",
    ])
    page_break(doc)

    title(doc, "4. THIẾT LẬP CÀI ĐẶT")
    para(doc, "Dự án gốc chưa có go.mod nên bước đầu tiên là khởi tạo Go module và tải dependency logger. Cache Go được đặt trong workspace khi chạy trong môi trường hạn chế quyền ghi.")
    code(doc, "git clone https://github.com/sdomino/scribble.git scribble-distributed\ncd scribble-distributed\ngo mod init github.com/sdomino/scribble\ngo mod tidy\ngo test ./...")
    table(doc, ["Yêu cầu", "Khuyến nghị"], [
        ("Go", "Phiên bản hỗ trợ Go modules"),
        ("Git", "Dùng để clone, branch, commit và push"),
        ("Cổng demo", "8081, 8082, 8083"),
        ("Hệ điều hành", "Windows/Linux/macOS"),
    ], [1.8, 4.5])
    page_break(doc)

    title(doc, "5. KẾT QUẢ BASELINE")
    para(doc, "Sau khi bổ sung go.mod và go.sum, toàn bộ kiểm thử gốc chạy thành công. Đây là mốc baseline để xác nhận phần mở rộng không làm hỏng hành vi hiện hữu.")
    code(doc, "ok   github.com/sdomino/scribble\n?    github.com/sdomino/scribble/example [no test files]")
    bullets(doc, [
        "TestNew: tạo và tái sử dụng thư mục database.",
        "TestWriteAndRead: ghi và đọc resource.",
        "TestReadall: đọc toàn bộ collection.",
        "TestDelete/TestDeleteall: xóa resource hoặc collection.",
    ])
    para(doc, "Nguyên tắc phát triển là giữ API cũ hoạt động nguyên vẹn. Khi Options không cấu hình Peers, Write và Delete chỉ thực hiện local như trước.")
    page_break(doc)

    title(doc, "6. YÊU CẦU CHO PHẦN MỞ RỘNG PHÂN TÁN")
    table(doc, ["Mã", "Yêu cầu", "Tiêu chí kiểm tra"], [
        ("R1", "Replication thao tác", "Write/Delete xuất hiện ở peer"),
        ("R2", "Không tạo vòng lặp", "Peer apply local, không forward tiếp"),
        ("R3", "Bootstrap node", "Node mới lấy được dữ liệu snapshot"),
        ("R4", "Tương thích API cũ", "Test gốc vẫn xanh"),
        ("R5", "Dễ thực nghiệm", "Có demo HTTP và lệnh PowerShell"),
    ], [0.7, 2.3, 3.3])
    para(doc, "Mô hình được chọn là primary-to-peers theo cấu hình tĩnh. Đây là lựa chọn vừa sức với phạm vi môn học: đủ thể hiện truyền thông phân tán, lỗi từng phần và đồng bộ trạng thái nhưng không đòi hỏi triển khai thuật toán consensus.")
    page_break(doc)

    title(doc, "7. TÍNH NĂNG MỚI 1: HTTP PEER REPLICATION")
    para(doc, "Driver được bổ sung NodeID, danh sách Peers và HTTPClient. Sau khi Write hoặc Delete local thành công, node tạo Operation và gửi POST tới endpoint /scribble/v1/operations của từng peer.")
    code(doc, "type Operation struct {\n    Type string\n    NodeID string\n    Collection string\n    Resource string\n    Data json.RawMessage\n}")
    bullets(doc, [
        "Operation write mang theo JSON data.",
        "Operation delete mang collection và resource.",
        "Peer gọi applyOperation để ghi/xóa local, không gọi lại replicate.",
        "Nếu một peer lỗi, node trả về lỗi replication có thông tin peer.",
    ])
    page_break(doc)

    title(doc, "8. LUỒNG HOẠT ĐỘNG REPLICATION")
    code(doc, "Client -> Node 1 Write local -> POST operation -> Node 2 apply local\n                         \\-> POST operation -> Node 3 apply local")
    para(doc, "Replication diễn ra đồng bộ theo thứ tự danh sách peer. Cách làm này đơn giản và giúp client biết ngay khi có peer không nhận được dữ liệu. Đổi lại, độ trễ ghi tăng theo số peer và trạng thái mạng.")
    table(doc, ["Trường hợp", "Hành vi"], [
        ("Không cấu hình peer", "Chỉ ghi local, tương thích phiên bản gốc"),
        ("Tất cả peer thành công", "Write/Delete trả nil"),
        ("Một peer lỗi", "Local vẫn thành công, trả lỗi replication"),
        ("Peer nhận operation", "Apply local, không phát tán tiếp"),
    ], [2.1, 4.2])
    page_break(doc)

    title(doc, "9. TÍNH NĂNG MỚI 2: SNAPSHOT VÀ SYNCFROM")
    para(doc, "Replication chỉ chuyển các thao tác mới. Khi một node mới tham gia hoặc một node bị mất dữ liệu, nó cần lấy trạng thái hiện có. CreateSnapshot duyệt thư mục database, thu thập các tệp JSON và trả về danh sách SnapshotRecord.")
    code(doc, "GET /scribble/v1/snapshot\n\nnewNode.SyncFrom(\"http://localhost:8081\")")
    bullets(doc, [
        "Snapshot chứa NodeID và toàn bộ record hợp lệ.",
        "SyncFrom tải snapshot qua HTTP và merge vào node đích.",
        "Việc apply snapshot không kích hoạt replication để tránh bão lưu lượng.",
        "Phù hợp cho bootstrap và sửa chữa thủ công trong demo.",
    ])
    page_break(doc)

    title(doc, "10. THIẾT KẾ ENDPOINT VÀ DỮ LIỆU")
    table(doc, ["Endpoint", "Method", "Chức năng", "Mã thành công"], [
        ("/scribble/v1/operations", "POST", "Nhận write/delete từ peer", "204"),
        ("/scribble/v1/snapshot", "GET", "Xuất snapshot hiện tại", "200"),
        ("/notes/{id}", "PUT/GET/DELETE", "API demo ứng dụng", "204/200"),
    ], [2.0, 0.8, 2.6, 0.9])
    para(doc, "Body được giới hạn kích thước khi decode để giảm rủi ro request quá lớn. HTTP client mặc định có timeout 3 giây. Trong hệ thống thực tế cần bổ sung xác thực giữa node, TLS, retry có backoff và idempotency key.")
    code(doc, '{"type":"write","node_id":"node-1","collection":"notes","resource":"demo","data":{"title":"Distributed Scribble"}}')
    page_break(doc)

    title(doc, "11. CHƯƠNG TRÌNH DEMO NHIỀU NODE")
    para(doc, "Thư mục example/distributed cung cấp một tiến trình node hoàn chỉnh. Mỗi tiến trình nhận địa chỉ listen, thư mục dữ liệu, NodeID, danh sách peer và peer dùng để bootstrap.")
    code(doc, "go run ./example/distributed -id node-2 -addr :8082 -data ./demo/node-2\ngo run ./example/distributed -id node-1 -addr :8081 -data ./demo/node-1 -peers http://localhost:8082")
    code(doc, "Invoke-RestMethod -Method Put -Uri http://localhost:8081/notes/demo -ContentType application/json -Body '{\"title\":\"Distributed Scribble\",\"body\":\"replicated\"}'\nInvoke-RestMethod -Method Get -Uri http://localhost:8082/notes/demo")
    para(doc, "Kết quả mong đợi: request PUT gửi vào node-1, tệp xuất hiện ở cả node-1 và node-2; request GET trực tiếp node-2 trả về cùng nội dung.")
    page_break(doc)

    title(doc, "12. KIỂM THỬ TỰ ĐỘNG")
    para(doc, "Kiểm thử mới sử dụng httptest.NewServer để tạo peer HTTP trong bộ nhớ. Cách này không phụ thuộc cổng mạng cố định và phù hợp chạy trong CI.")
    table(doc, ["Test", "Mục tiêu"], [
        ("TestReplicationWriteAndDelete", "Xác nhận write và delete được sao chép"),
        ("TestSyncFromSnapshot", "Xác nhận node đích lấy được record từ snapshot"),
        ("Toàn bộ test gốc", "Phát hiện regression API local"),
    ], [2.8, 3.5])
    code(doc, "go test ./...\n\nok   github.com/sdomino/scribble\n?    github.com/sdomino/scribble/example/distributed [no test files]")
    page_break(doc)

    title(doc, "13. KẾT QUẢ VÀ ĐÁNH GIÁ")
    bullets(doc, [
        "Tất cả test chạy thành công sau khi thêm hai tính năng.",
        "Write/Delete vẫn hoạt động như cũ khi không có peer.",
        "Replication tạo bản sao ở node phụ và xóa đồng bộ.",
        "SyncFrom giúp node mới nhận dữ liệu có sẵn.",
    ])
    table(doc, ["Tiêu chí", "Đạt", "Nhận xét"], [
        ("Tính đúng đắn chức năng", "Có", "Được kiểm chứng bằng automated tests"),
        ("Khả năng quan sát", "Có", "Demo HTTP và dữ liệu theo thư mục"),
        ("Khả năng mở rộng", "Một phần", "Gửi tuần tự; phù hợp demo nhỏ"),
        ("Nhất quán mạnh", "Không", "Không có consensus/quorum"),
    ], [2.0, 0.8, 3.5])
    page_break(doc)

    title(doc, "14. GIỚI HẠN VÀ HƯỚNG PHÁT TRIỂN")
    para(doc, "Mô hình hiện tại minh họa replication nhưng chưa giải quyết đầy đủ các bài toán của hệ phân tán. Nếu hai node cùng ghi một resource, giá trị đến sau sẽ ghi đè mà không có version hoặc phát hiện xung đột.")
    bullets(doc, [
        "Bổ sung operation ID và hàng đợi retry bền vững.",
        "Dùng version/vector clock để phát hiện ghi đồng thời.",
        "Thêm quorum read/write hoặc thuật toán Raft cho leader election.",
        "Snapshot nhất quán hơn bằng khóa toàn cục hoặc copy-on-write.",
        "Thêm authentication, TLS, metrics và health check.",
    ])
    para(doc, "Những hướng trên có thể được chia thành các milestone tiếp theo, nhưng vượt ngoài phạm vi hai tính năng yêu cầu của bài tập.")
    page_break(doc)

    title(doc, "15. QUẢN LÝ MÃ NGUỒN TRÊN GITHUB")
    para(doc, "Repo cần được tạo dưới tài khoản hoặc organization của nhóm, sau đó push toàn bộ lịch sử phát triển. Không squash thành một commit cuối nếu giảng viên cần đánh giá quá trình.")
    code(doc, "git remote rename origin upstream\ngit remote add origin https://github.com/<group>/scribble-distributed.git\ngit push -u origin master")
    table(doc, ["Commit mẫu đã tạo", "Nội dung"], [
        ("Initialize Go module for coursework", "Thiết lập module và dependency"),
        ("Add peer replication and snapshot sync", "Hai tính năng phân tán + test"),
        ("Document and demonstrate distributed nodes", "README và demo nhiều node"),
        ("Add coursework report and presentation", "Bộ tài liệu nộp bài"),
    ], [2.8, 3.5])
    page_break(doc)

    title(doc, "16. PHÂN CÔNG VÀ ĐÓNG GÓP THÀNH VIÊN")
    para(doc, "Thay các ô bên dưới bằng thông tin thật. Mỗi thành viên nên có commit mang tên/email đúng và pull request hoặc issue tương ứng.")
    table(doc, ["Thành viên", "Phần việc", "Minh chứng GitHub"], [
        ("[Họ tên 1 - MSSV]", "Phân tích Scribble, setup module, baseline tests", "[Commit/PR link]"),
        ("[Họ tên 2 - MSSV]", "Replication HTTP, operation handler, tests", "[Commit/PR link]"),
        ("[Họ tên 3 - MSSV]", "Snapshot/SyncFrom, demo nhiều node", "[Commit/PR link]"),
        ("[Họ tên 4 - MSSV]", "Báo cáo, slides, thực nghiệm và QA", "[Commit/PR link]"),
    ], [1.7, 3.0, 1.6])
    bullets(doc, [
        "Khuyến nghị: mỗi người tạo branch feature riêng và mở pull request.",
        "Ghi issue cho từng tính năng, gắn assignee và milestone.",
        "Chụp trang Insights/Contributors và Network graph để đưa vào phần phụ lục nếu cần.",
    ])
    page_break(doc)

    title(doc, "17. KẾT LUẬN")
    para(doc, "Bài tập đã chuyển Scribble từ một thư viện lưu JSON đơn node thành prototype có khả năng trao đổi dữ liệu giữa nhiều node. Replication HTTP giải quyết việc chuyển tiếp thao tác mới; snapshot/SyncFrom giải quyết bootstrap và phục hồi dữ liệu hiện có.")
    para(doc, "Thiết kế vẫn giữ được tinh thần tối giản của dự án gốc, có test tự động và chương trình demo để tái tạo kết quả. Quan trọng hơn, phần mở rộng làm rõ các vấn đề thực tế của hệ phân tán: lỗi từng phần, độ trễ mạng, vòng lặp replication, xung đột và tính nhất quán.")
    bullets(doc, [
        "Hoàn thành hai tính năng mới liên quan xử lý phân tán.",
        "Có mã nguồn, kiểm thử, hướng dẫn chạy, báo cáo và slides.",
        "Có lịch sử commit theo từng giai đoạn để push lên GitHub nhóm.",
    ])

    title(doc, "TÀI LIỆU THAM KHẢO")
    bullets(doc, [
        "Scribble source repository: https://github.com/sdomino/scribble",
        "The Go Programming Language documentation: https://go.dev/doc/",
        "Go net/http package: https://pkg.go.dev/net/http",
        "Go testing and httptest packages: https://pkg.go.dev/testing và https://pkg.go.dev/net/http/httptest",
        "Martin Kleppmann, Designing Data-Intensive Applications, O'Reilly Media.",
    ])
    para(doc, "Ghi chú: báo cáo mô tả prototype phục vụ học tập. Các nhận định về giới hạn và hướng phát triển được rút ra từ thiết kế và kết quả thực nghiệm của nhóm.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
